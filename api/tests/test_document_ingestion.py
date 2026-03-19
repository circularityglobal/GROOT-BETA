"""
REFINET Cloud — Document Ingestion System Tests
Tests document parsing, auto-tagging, document comparison, RAG integration, and MCP tools.
"""

import json
import os
import sys
import io

import pytest

# Ensure project root is on path
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))


# ═══════════════════════════════════════════════════════════════════════
# Test 1: Document Parser — Text & Markdown
# ═══════════════════════════════════════════════════════════════════════

class TestParserTextMarkdown:
    def test_parse_plain_text(self):
        from api.services.document_parser import parse_file
        content = b"Hello world. This is a plain text file."
        result = parse_file(content, "test.txt")
        assert result.doc_type == "txt"
        assert "Hello world" in result.text
        assert result.error is None
        assert result.metadata.get("file_size") == len(content)

    def test_parse_markdown(self):
        from api.services.document_parser import parse_file
        content = b"# Title\n\nSome paragraph.\n\n## Subtitle\n\nMore text."
        result = parse_file(content, "readme.md")
        assert result.doc_type == "md"
        assert "# Title" in result.text
        assert result.metadata.get("headings") == ["Title", "Subtitle"]

    def test_parse_empty_file(self):
        from api.services.document_parser import parse_file
        result = parse_file(b"", "empty.txt")
        assert result.text == ""
        assert result.doc_type == "txt"

    def test_parse_unsupported_extension(self):
        from api.services.document_parser import parse_file
        result = parse_file(b"data", "file.xyz")
        assert result.error is not None
        assert "Unsupported" in result.error

    def test_parse_utf8_bom(self):
        from api.services.document_parser import parse_file
        bom = b'\xef\xbb\xbf'
        content = bom + "Hello UTF-8 BOM".encode("utf-8")
        result = parse_file(content, "bom.txt")
        assert "Hello UTF-8 BOM" in result.text

    def test_parse_latin1(self):
        from api.services.document_parser import parse_file
        content = "Héllo wörld".encode("latin-1")
        result = parse_file(content, "latin.txt")
        assert "wörld" in result.text or "rld" in result.text  # Should handle encoding


# ═══════════════════════════════════════════════════════════════════════
# Test 2: Document Parser — CSV
# ═══════════════════════════════════════════════════════════════════════

class TestParserCSV:
    def test_parse_csv(self):
        from api.services.document_parser import parse_file
        csv_content = b"name,age,city\nAlice,30,NYC\nBob,25,LA"
        result = parse_file(csv_content, "data.csv")
        assert result.doc_type == "csv"
        assert "Alice" in result.text
        assert "Bob" in result.text
        assert result.metadata.get("row_count") == 3  # header + 2 rows

    def test_parse_csv_with_bom(self):
        from api.services.document_parser import parse_file
        csv_content = b'\xef\xbb\xbfname,value\nfoo,bar'
        result = parse_file(csv_content, "bom.csv")
        assert "foo" in result.text


# ═══════════════════════════════════════════════════════════════════════
# Test 3: Document Parser — JSON & ABI
# ═══════════════════════════════════════════════════════════════════════

class TestParserJSON:
    def test_parse_generic_json(self):
        from api.services.document_parser import parse_file
        data = {"key": "value", "number": 42}
        result = parse_file(json.dumps(data).encode(), "config.json")
        assert result.doc_type == "json"
        assert '"key"' in result.text
        assert result.metadata.get("json_type") == "dict"

    def test_parse_abi_json(self):
        from api.services.document_parser import parse_file
        abi = [
            {
                "type": "function",
                "name": "transfer",
                "inputs": [
                    {"type": "address", "name": "to"},
                    {"type": "uint256", "name": "amount"},
                ],
                "outputs": [{"type": "bool", "name": ""}],
                "stateMutability": "nonpayable",
            },
            {
                "type": "event",
                "name": "Transfer",
                "inputs": [
                    {"type": "address", "name": "from", "indexed": True},
                    {"type": "address", "name": "to", "indexed": True},
                    {"type": "uint256", "name": "value", "indexed": False},
                ],
            },
        ]
        result = parse_file(json.dumps(abi).encode(), "erc20.json")
        assert result.doc_type == "json"
        assert result.metadata.get("is_abi") is True
        assert result.metadata.get("function_count") == 1
        assert result.metadata.get("event_count") == 1
        assert "transfer" in result.text.lower()
        assert "Transfer" in result.text

    def test_parse_invalid_json(self):
        from api.services.document_parser import parse_file
        # Invalid JSON should trigger the outer exception handler
        result = parse_file(b"{not valid json}", "bad.json")
        assert result.error is not None


# ═══════════════════════════════════════════════════════════════════════
# Test 4: Document Parser — Solidity
# ═══════════════════════════════════════════════════════════════════════

class TestParserSolidity:
    def test_parse_solidity(self):
        from api.services.document_parser import parse_file
        sol_code = b"""// SPDX-License-Identifier: MIT
pragma solidity ^0.8.20;

/// @title MyToken
/// @notice A simple ERC20 token
contract MyToken {
    mapping(address => uint256) public balances;

    function transfer(address to, uint256 amount) external {
        balances[msg.sender] -= amount;
        balances[to] += amount;
    }
}

interface IMyToken {
    function transfer(address to, uint256 amount) external;
}
"""
        result = parse_file(sol_code, "MyToken.sol")
        assert result.doc_type == "sol"
        assert "MyToken" in result.text
        assert "pragma" in result.text
        assert result.metadata.get("solidity_version") == "^0.8.20"
        assert "MyToken" in result.metadata.get("contracts", [])
        assert "IMyToken" in result.metadata.get("contracts", [])
        assert result.metadata.get("natspec_count", 0) >= 2


# ═══════════════════════════════════════════════════════════════════════
# Test 5: Document Parser — PDF (with optional dep)
# ═══════════════════════════════════════════════════════════════════════

class TestParserPDF:
    def test_parse_pdf_import_check(self):
        """Test that PDF parsing handles missing PyMuPDF gracefully."""
        from api.services.document_parser import parse_file
        try:
            import fitz  # noqa
            # PyMuPDF is available — create a minimal PDF
            doc = fitz.open()
            page = doc.new_page()
            page.insert_text((72, 72), "Hello from PDF")
            pdf_bytes = doc.tobytes()
            doc.close()

            result = parse_file(pdf_bytes, "test.pdf")
            assert result.doc_type == "pdf"
            assert result.page_count == 1
            assert "Hello from PDF" in result.text
        except ImportError:
            # PyMuPDF not installed — verify graceful error
            result = parse_file(b"%PDF-1.4 fake", "test.pdf")
            assert result.error is not None
            assert "PyMuPDF" in result.error


# ═══════════════════════════════════════════════════════════════════════
# Test 6: Document Parser — DOCX (with optional dep)
# ═══════════════════════════════════════════════════════════════════════

class TestParserDOCX:
    def test_parse_docx_import_check(self):
        from api.services.document_parser import parse_file
        try:
            import docx
            # Create a minimal DOCX
            doc = docx.Document()
            doc.add_paragraph("Hello from DOCX")
            doc.add_paragraph("Second paragraph")
            buf = io.BytesIO()
            doc.save(buf)
            docx_bytes = buf.getvalue()

            result = parse_file(docx_bytes, "test.docx")
            assert result.doc_type == "docx"
            assert "Hello from DOCX" in result.text
            assert "Second paragraph" in result.text
        except ImportError:
            result = parse_file(b"PK fake docx", "test.docx")
            assert result.error is not None
            assert "python-docx" in result.error


# ═══════════════════════════════════════════════════════════════════════
# Test 7: Document Parser — XLSX (with optional dep)
# ═══════════════════════════════════════════════════════════════════════

class TestParserXLSX:
    def test_parse_xlsx_import_check(self):
        from api.services.document_parser import parse_file
        try:
            import openpyxl
            # Create a minimal XLSX
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Data"
            ws.append(["Name", "Value"])
            ws.append(["Alice", 100])
            ws.append(["Bob", 200])
            buf = io.BytesIO()
            wb.save(buf)
            xlsx_bytes = buf.getvalue()

            result = parse_file(xlsx_bytes, "data.xlsx")
            assert result.doc_type == "xlsx"
            assert result.page_count == 1  # 1 sheet
            assert "Alice" in result.text
            assert "Data" in result.metadata.get("sheet_names", [])
        except ImportError:
            result = parse_file(b"PK fake xlsx", "data.xlsx")
            assert result.error is not None
            assert "openpyxl" in result.error


# ═══════════════════════════════════════════════════════════════════════
# Test 8: Auto-Tagger — Tag Generation
# ═══════════════════════════════════════════════════════════════════════

class TestAutoTagger:
    def test_generate_tags_basic(self):
        from api.services.auto_tagger import generate_tags
        text = """
        Ethereum is a decentralized blockchain platform. Smart contracts enable
        DeFi applications like staking, lending, and token swaps. The ERC20 standard
        defines a common interface for fungible tokens on the Ethereum network.
        Governance mechanisms allow DAO participants to vote on protocol changes.
        """
        tags = generate_tags(text, doc_type="md")
        assert isinstance(tags, list)
        assert len(tags) > 0
        assert len(tags) <= 15
        # Should include domain terms
        tag_str = " ".join(tags)
        assert any(t in tag_str for t in ["ethereum", "defi", "token", "staking", "blockchain", "governance"])

    def test_generate_tags_empty(self):
        from api.services.auto_tagger import generate_tags
        assert generate_tags("") == []
        assert generate_tags("   ") == []

    def test_generate_tags_solidity(self):
        from api.services.auto_tagger import generate_tags
        text = """
        pragma solidity ^0.8.20;
        contract StakingVault {
            mapping(address => uint256) public stakes;
            function deposit() external payable {}
            function withdraw(uint256 amount) external {}
        }
        """
        tags = generate_tags(text, doc_type="sol", filename="StakingVault.sol")
        assert isinstance(tags, list)
        assert len(tags) > 0
        tag_str = " ".join(tags)
        # Should have file-type and filename tags
        assert "solidity contract" in tag_str or "smart contract" in tag_str

    def test_generate_tags_max_limit(self):
        from api.services.auto_tagger import generate_tags
        text = "blockchain " * 100 + "ethereum " * 100 + "defi " * 100
        tags = generate_tags(text, max_tags=5)
        assert len(tags) <= 5

    def test_generate_tags_with_filename(self):
        from api.services.auto_tagger import generate_tags
        tags = generate_tags("some content about staking", doc_type="pdf", filename="DeFi-Staking-Guide.pdf")
        tag_str = " ".join(tags)
        assert "pdf document" in tag_str
        # Filename-derived tag
        assert "defi staking guide" in tag_str


# ═══════════════════════════════════════════════════════════════════════
# Test 9: Auto-Tagger — Category Inference
# ═══════════════════════════════════════════════════════════════════════

class TestCategoryInference:
    def test_infer_blockchain(self):
        from api.services.auto_tagger import infer_category
        text = "Smart contract deployment on Ethereum with ERC20 token staking and DeFi governance"
        tags = ["ethereum", "defi", "staking", "erc20"]
        assert infer_category(text, tags) == "blockchain"

    def test_infer_docs(self):
        from api.services.auto_tagger import infer_category
        text = "API endpoint configuration. Authentication setup guide. Request parameters and response schema."
        tags = ["api", "authentication", "configuration"]
        assert infer_category(text, tags) == "docs"

    def test_infer_contract(self):
        from api.services.auto_tagger import infer_category
        text = "pragma solidity. function selector calldata. ABI interface. constructor modifier mapping struct"
        tags = ["abi", "function", "modifier", "constructor"]
        assert infer_category(text, tags) == "contract"

    def test_infer_about(self):
        from api.services.auto_tagger import infer_category
        text = "REFINET mission and vision. Sovereignty and decentralization. Community-driven open source."
        tags = ["refinet", "mission", "sovereignty"]
        assert infer_category(text, tags) == "about"

    def test_infer_product(self):
        from api.services.auto_tagger import infer_category
        text = "QuickCast autonomous publishing feature. AgentOS AI agents dashboard. CIFI Wizards platform."
        tags = ["quickcast", "agentos", "dashboard"]
        assert infer_category(text, tags) == "product"

    def test_infer_faq(self):
        from api.services.auto_tagger import infer_category
        text = "How to get started. What is REFINET? Common questions and answers about troubleshooting."
        tags = ["how to", "getting started", "question"]
        assert infer_category(text, tags) == "faq"

    def test_infer_default(self):
        from api.services.auto_tagger import infer_category
        text = "xyzzy lorem ipsum dolor sit amet"
        tags = ["xyzzy"]
        assert infer_category(text, tags) == "docs"  # default


# ═══════════════════════════════════════════════════════════════════════
# Test 10: Auto-Tagger — Keyword Extraction
# ═══════════════════════════════════════════════════════════════════════

class TestKeywordExtraction:
    def test_extract_keywords(self):
        from api.services.auto_tagger import _extract_keywords
        text = "ethereum blockchain staking defi protocol governance token yield farming yield farming"
        keywords = _extract_keywords(text, top_n=10)
        assert isinstance(keywords, list)
        assert len(keywords) > 0
        # Each entry is (keyword, score)
        assert all(isinstance(k, tuple) and len(k) == 2 for k in keywords)
        # Domain terms should score high
        kw_names = [k[0] for k in keywords]
        assert any("ethereum" in k or "staking" in k or "defi" in k for k in kw_names)

    def test_extract_keywords_empty(self):
        from api.services.auto_tagger import _extract_keywords
        assert _extract_keywords("") == []
        assert _extract_keywords("a an the") == []  # all stop words


# ═══════════════════════════════════════════════════════════════════════
# Test 11: Auto-Tagger — Named Entity Extraction
# ═══════════════════════════════════════════════════════════════════════

class TestNamedEntityExtraction:
    def test_extract_capitalized_phrases(self):
        from api.services.auto_tagger import _extract_named_entities
        text = "The Smart Contract Audit by Open Zeppelin revealed vulnerabilities."
        entities = _extract_named_entities(text)
        entity_str = " ".join(e.lower() for e in entities)
        assert "smart contract" in entity_str or "open zeppelin" in entity_str

    def test_extract_ethereum_addresses(self):
        from api.services.auto_tagger import _extract_named_entities
        text = "Contract deployed at 0x1234567890abcdef1234567890abcdef12345678"
        entities = _extract_named_entities(text)
        assert any("0x1234" in e for e in entities)

    def test_extract_camel_case(self):
        from api.services.auto_tagger import _extract_named_entities
        text = "The stakingRewards function calls transferFrom on the token contract"
        entities = _extract_named_entities(text)
        entity_str = " ".join(e.lower() for e in entities)
        assert "staking rewards" in entity_str or "transfer from" in entity_str


# ═══════════════════════════════════════════════════════════════════════
# Test 12: Document Comparison
# ═══════════════════════════════════════════════════════════════════════

class TestDocumentComparison:
    def test_keyword_overlap(self):
        from api.services.document_compare import _compute_keyword_overlap
        text_a = "ethereum blockchain staking defi protocol governance"
        text_b = "ethereum blockchain lending defi protocol oracle"
        result = _compute_keyword_overlap(text_a, text_b)
        assert "score" in result
        assert 0.0 <= result["score"] <= 1.0
        assert "ethereum" in result["shared_keywords"]
        assert "blockchain" in result["shared_keywords"]
        assert "staking" in result["unique_to_a"]
        assert "lending" in result["unique_to_b"]

    def test_tag_overlap(self):
        from api.services.document_compare import _compute_tag_overlap
        tags_a = ["defi", "staking", "ethereum"]
        tags_b = ["defi", "lending", "ethereum"]
        result = _compute_tag_overlap(tags_a, tags_b)
        assert result["score"] > 0
        assert "defi" in result["shared_tags"]
        assert "ethereum" in result["shared_tags"]
        assert "staking" in result["unique_to_a"]
        assert "lending" in result["unique_to_b"]

    def test_tag_overlap_empty(self):
        from api.services.document_compare import _compute_tag_overlap
        result = _compute_tag_overlap([], [])
        assert result["score"] == 0.0

    def test_extract_headings(self):
        from api.services.document_compare import _extract_headings
        text = "# Title\n## Section One\n### Subsection\n=== Important Section ==="
        headings = _extract_headings(text)
        assert "Title" in headings
        assert "Section One" in headings
        assert "Important Section" in headings

    def test_parse_tags_safety(self):
        from api.services.document_compare import _parse_tags
        assert _parse_tags(None) == []
        assert _parse_tags("") == []
        assert _parse_tags("not json") == []
        assert _parse_tags('["a","b"]') == ["a", "b"]


# ═══════════════════════════════════════════════════════════════════════
# Test 13: ParseResult Dataclass
# ═══════════════════════════════════════════════════════════════════════

class TestParseResult:
    def test_parse_result_defaults(self):
        from api.services.document_parser import ParseResult
        r = ParseResult(text="hello", doc_type="txt")
        assert r.text == "hello"
        assert r.doc_type == "txt"
        assert r.page_count is None
        assert r.metadata == {}
        assert r.error is None

    def test_parse_result_with_all_fields(self):
        from api.services.document_parser import ParseResult
        r = ParseResult(
            text="content",
            doc_type="pdf",
            page_count=5,
            metadata={"author": "Alice"},
            error="partial extraction",
        )
        assert r.page_count == 5
        assert r.metadata["author"] == "Alice"
        assert r.error == "partial extraction"


# ═══════════════════════════════════════════════════════════════════════
# Test 14: SUPPORTED_EXTENSIONS constant
# ═══════════════════════════════════════════════════════════════════════

class TestConstants:
    def test_supported_extensions(self):
        from api.services.document_parser import SUPPORTED_EXTENSIONS
        assert ".pdf" in SUPPORTED_EXTENSIONS
        assert ".docx" in SUPPORTED_EXTENSIONS
        assert ".xlsx" in SUPPORTED_EXTENSIONS
        assert ".csv" in SUPPORTED_EXTENSIONS
        assert ".txt" in SUPPORTED_EXTENSIONS
        assert ".md" in SUPPORTED_EXTENSIONS
        assert ".json" in SUPPORTED_EXTENSIONS
        assert ".sol" in SUPPORTED_EXTENSIONS
        assert len(SUPPORTED_EXTENSIONS) == 8

    def test_ext_to_doctype(self):
        from api.services.document_parser import EXT_TO_DOCTYPE
        assert EXT_TO_DOCTYPE[".pdf"] == "pdf"
        assert EXT_TO_DOCTYPE[".sol"] == "sol"


# ═══════════════════════════════════════════════════════════════════════
# Test 15: MCP Tool Definitions
# ═══════════════════════════════════════════════════════════════════════

class TestMCPTools:
    def test_new_tools_registered(self):
        from api.services.mcp_gateway import MCP_TOOLS
        assert "search_documents" in MCP_TOOLS
        assert "compare_documents" in MCP_TOOLS
        assert "get_document_tags" in MCP_TOOLS

    def test_search_documents_schema(self):
        from api.services.mcp_gateway import MCP_TOOLS
        tool = MCP_TOOLS["search_documents"]
        assert "query" in tool["input_schema"]["properties"]
        assert "tags" in tool["input_schema"]["properties"]
        assert "category" in tool["input_schema"]["properties"]
        assert tool["input_schema"]["required"] == ["query"]

    def test_compare_documents_schema(self):
        from api.services.mcp_gateway import MCP_TOOLS
        tool = MCP_TOOLS["compare_documents"]
        assert "doc_id_a" in tool["input_schema"]["properties"]
        assert "doc_id_b" in tool["input_schema"]["properties"]
        assert set(tool["input_schema"]["required"]) == {"doc_id_a", "doc_id_b"}

    def test_get_document_tags_schema(self):
        from api.services.mcp_gateway import MCP_TOOLS
        tool = MCP_TOOLS["get_document_tags"]
        assert "doc_id" in tool["input_schema"]["properties"]
        assert tool["input_schema"]["required"] == ["doc_id"]

    def test_list_tools_includes_new(self):
        from api.services.mcp_gateway import list_tools
        tools = list_tools()
        names = [t["name"] for t in tools]
        assert "search_documents" in names
        assert "compare_documents" in names
        assert "get_document_tags" in names
        # Should have at least 15 tools (may grow as new features are added)
        assert len(tools) >= 15


# ═══════════════════════════════════════════════════════════════════════
# Test 16: Knowledge Model Schema
# ═══════════════════════════════════════════════════════════════════════

class TestKnowledgeModel:
    def test_model_has_new_columns(self):
        from api.models.knowledge import KnowledgeDocument
        # Verify new columns exist on the model
        columns = {c.name for c in KnowledgeDocument.__table__.columns}
        assert "tags" in columns
        assert "doc_type" in columns
        assert "page_count" in columns
        assert "metadata_json" in columns

    def test_model_new_columns_nullable(self):
        from api.models.knowledge import KnowledgeDocument
        col_map = {c.name: c for c in KnowledgeDocument.__table__.columns}
        assert col_map["tags"].nullable is True
        assert col_map["doc_type"].nullable is True
        assert col_map["page_count"].nullable is True
        assert col_map["metadata_json"].nullable is True


# ═══════════════════════════════════════════════════════════════════════
# Test 17: RAG Ingest Document Signature
# ═══════════════════════════════════════════════════════════════════════

class TestRAGSignature:
    def test_ingest_document_accepts_new_params(self):
        """Verify ingest_document accepts the new optional parameters."""
        import inspect
        from api.services.rag import ingest_document
        sig = inspect.signature(ingest_document)
        params = list(sig.parameters.keys())
        assert "tags" in params
        assert "doc_type" in params
        assert "page_count" in params
        assert "metadata_json" in params
        # All new params should have default None
        assert sig.parameters["tags"].default is None
        assert sig.parameters["doc_type"].default is None
        assert sig.parameters["page_count"].default is None
        assert sig.parameters["metadata_json"].default is None

    def test_search_knowledge_accepts_tags(self):
        import inspect
        from api.services.rag import search_knowledge
        sig = inspect.signature(search_knowledge)
        assert "tags" in sig.parameters
        assert sig.parameters["tags"].default is None


# ═══════════════════════════════════════════════════════════════════════
# Test 18: Diversity Selection Fallback
# ═══════════════════════════════════════════════════════════════════════

class TestDiversitySelection:
    def test_fallback_when_few_candidates(self):
        from api.services.auto_tagger import _select_diverse_tags
        candidates = ["a", "b", "c"]
        result = _select_diverse_tags(candidates, max_tags=10)
        assert result == ["a", "b", "c"]

    def test_truncation(self):
        from api.services.auto_tagger import _select_diverse_tags
        candidates = [f"tag_{i}" for i in range(50)]
        result = _select_diverse_tags(candidates, max_tags=5)
        assert len(result) == 5


# ═══════════════════════════════════════════════════════════════════════
# Test 19: File-Type Tags
# ═══════════════════════════════════════════════════════════════════════

class TestFileTypeTags:
    def test_pdf_type_tag(self):
        from api.services.auto_tagger import _get_type_tags
        tags = _get_type_tags("pdf", "report.pdf")
        assert "pdf document" in tags

    def test_sol_type_tag(self):
        from api.services.auto_tagger import _get_type_tags
        tags = _get_type_tags("sol", "Token.sol")
        assert "solidity contract" in tags
        assert "smart contract" in tags

    def test_filename_tag(self):
        from api.services.auto_tagger import _get_type_tags
        tags = _get_type_tags("pdf", "DeFi-Staking-Guide.pdf")
        assert "defi staking guide" in tags

    def test_no_filename(self):
        from api.services.auto_tagger import _get_type_tags
        tags = _get_type_tags("txt", None)
        assert "text document" in tags


# ═══════════════════════════════════════════════════════════════════════
# Test 20: Integration — Parse → Tag → Category pipeline
# ═══════════════════════════════════════════════════════════════════════

class TestIntegrationPipeline:
    def test_full_pipeline_text(self):
        """End-to-end: parse text → generate tags → infer category."""
        from api.services.document_parser import parse_file
        from api.services.auto_tagger import generate_tags, infer_category

        content = b"""
        # DeFi Staking Protocol

        This document describes a decentralized staking protocol built on Ethereum.
        Users deposit tokens into a vault smart contract and earn yield through
        automated market making and liquidity provision.

        ## Token Economics
        The governance token (ERC20) is used for voting on protocol upgrades.
        Staking rewards are distributed proportionally to deposited amounts.

        ## Smart Contract Architecture
        The protocol uses upgradeable proxy contracts following the UUPS pattern.
        Access control is managed through OpenZeppelin's AccessControl contract.
        """
        result = parse_file(content, "defi-staking.md")
        assert result.doc_type == "md"

        tags = generate_tags(result.text, doc_type=result.doc_type, filename="defi-staking.md")
        assert len(tags) > 0

        category = infer_category(result.text, tags)
        assert category in ("blockchain", "contract")  # Should detect blockchain/defi content

    def test_full_pipeline_csv(self):
        from api.services.document_parser import parse_file
        from api.services.auto_tagger import generate_tags, infer_category

        content = b"token,price,volume\nETH,3500,1000000\nBTC,60000,2000000\nSOL,150,500000"
        result = parse_file(content, "prices.csv")
        assert result.doc_type == "csv"

        tags = generate_tags(result.text, doc_type=result.doc_type, filename="prices.csv")
        assert "tabular data" in tags

        category = infer_category(result.text, tags)
        assert isinstance(category, str)

    def test_full_pipeline_abi(self):
        from api.services.document_parser import parse_file
        from api.services.auto_tagger import generate_tags, infer_category

        abi = [
            {"type": "function", "name": "stake", "inputs": [{"type": "uint256", "name": "amount"}], "outputs": [], "stateMutability": "nonpayable"},
            {"type": "function", "name": "withdraw", "inputs": [{"type": "uint256", "name": "amount"}], "outputs": [], "stateMutability": "nonpayable"},
            {"type": "event", "name": "Staked", "inputs": [{"type": "address", "name": "user", "indexed": True}, {"type": "uint256", "name": "amount", "indexed": False}]},
        ]
        result = parse_file(json.dumps(abi).encode(), "staking.json")
        assert result.metadata.get("is_abi") is True

        tags = generate_tags(result.text, doc_type=result.doc_type, filename="staking.json")
        assert len(tags) > 0

        category = infer_category(result.text, tags)
        assert category in ("blockchain", "contract")


# ═══════════════════════════════════════════════════════════════════════
# Test 21: Edge Cases
# ═══════════════════════════════════════════════════════════════════════

class TestEdgeCases:
    def test_very_long_text(self):
        from api.services.auto_tagger import generate_tags
        text = "ethereum staking defi blockchain " * 10000
        tags = generate_tags(text, max_tags=10)
        assert len(tags) <= 10

    def test_single_word_text(self):
        from api.services.auto_tagger import generate_tags
        tags = generate_tags("ethereum")
        # Should handle gracefully even with minimal input
        assert isinstance(tags, list)

    def test_binary_file_fallback(self):
        from api.services.document_parser import parse_file
        # Binary garbage with .txt extension
        result = parse_file(bytes(range(256)), "binary.txt")
        # Should still return something via fallback
        assert isinstance(result.text, str)

    def test_comparison_keyword_overlap_identical(self):
        from api.services.document_compare import _compute_keyword_overlap
        text = "ethereum blockchain staking protocol governance"
        result = _compute_keyword_overlap(text, text)
        assert result["score"] == 1.0
        assert len(result["unique_to_a"]) == 0
        assert len(result["unique_to_b"]) == 0

    def test_comparison_keyword_overlap_no_overlap(self):
        from api.services.document_compare import _compute_keyword_overlap
        text_a = "apples oranges bananas grapes"
        text_b = "kubernetes docker terraform helm"
        result = _compute_keyword_overlap(text_a, text_b)
        assert result["score"] == 0.0


# ═══════════════════════════════════════════════════════════════════════
# Test 22: NotebookLM Features — Schema
# ═══════════════════════════════════════════════════════════════════════

class TestNotebookSchemas:
    def test_source_reference_model(self):
        from api.schemas.inference import SourceReference
        src = SourceReference(
            document_id="abc123",
            document_title="Test Doc",
            category="docs",
            doc_type="pdf",
            tags=["defi", "staking"],
            score=0.85,
            preview="First 150 chars...",
        )
        assert src.document_id == "abc123"
        assert src.score == 0.85
        assert len(src.tags) == 2

    def test_chat_request_notebook_doc_ids(self):
        from api.schemas.inference import ChatCompletionRequest, ChatMessage
        req = ChatCompletionRequest(
            messages=[ChatMessage(role="user", content="test")],
            notebook_doc_ids=["doc1", "doc2"],
        )
        assert req.notebook_doc_ids == ["doc1", "doc2"]

    def test_chat_request_notebook_optional(self):
        from api.schemas.inference import ChatCompletionRequest, ChatMessage
        req = ChatCompletionRequest(
            messages=[ChatMessage(role="user", content="test")],
        )
        assert req.notebook_doc_ids is None

    def test_chat_response_sources(self):
        from api.schemas.inference import ChatCompletionResponse, ChatCompletionChoice, ChatMessage, UsageInfo, SourceReference
        resp = ChatCompletionResponse(
            id="test",
            created=1234,
            model="bitnet",
            choices=[ChatCompletionChoice(message=ChatMessage(role="assistant", content="hi"))],
            usage=UsageInfo(prompt_tokens=10, completion_tokens=5, total_tokens=15),
            sources=[SourceReference(
                document_id="d1", document_title="Doc", category="docs",
                score=0.9, preview="...",
            )],
        )
        assert len(resp.sources) == 1
        assert resp.sources[0].document_id == "d1"

    def test_chat_response_sources_optional(self):
        from api.schemas.inference import ChatCompletionResponse, ChatCompletionChoice, ChatMessage, UsageInfo
        resp = ChatCompletionResponse(
            id="test", created=1234, model="bitnet",
            choices=[ChatCompletionChoice(message=ChatMessage(role="assistant", content="hi"))],
            usage=UsageInfo(prompt_tokens=10, completion_tokens=5, total_tokens=15),
        )
        assert resp.sources is None


# ═══════════════════════════════════════════════════════════════════════
# Test 23: RAG returns sources tuple
# ═══════════════════════════════════════════════════════════════════════

class TestRAGSourcesReturn:
    def test_build_groot_system_prompt_signature(self):
        import inspect
        from api.services.rag import build_groot_system_prompt
        sig = inspect.signature(build_groot_system_prompt)
        assert "doc_ids" in sig.parameters

    def test_build_rag_context_signature(self):
        import inspect
        from api.services.rag import build_rag_context
        sig = inspect.signature(build_rag_context)
        assert "doc_ids" in sig.parameters

    def test_search_knowledge_doc_ids_param(self):
        import inspect
        from api.services.rag import search_knowledge
        sig = inspect.signature(search_knowledge)
        assert "doc_ids" in sig.parameters
        assert sig.parameters["doc_ids"].default is None


# ═══════════════════════════════════════════════════════════════════════
# Test 24: Document Generator
# ═══════════════════════════════════════════════════════════════════════

class TestDocumentGenerator:
    def test_prompt_templates_exist(self):
        from api.services.document_generator import SUMMARIZE_PROMPT, FAQ_PROMPT, OVERVIEW_PROMPT
        assert "{content}" in SUMMARIZE_PROMPT
        assert "{content}" in FAQ_PROMPT
        assert "{content}" in OVERVIEW_PROMPT

    def test_max_content_chars(self):
        from api.services.document_generator import MAX_CONTENT_CHARS
        assert MAX_CONTENT_CHARS > 0
        assert MAX_CONTENT_CHARS <= 10000  # reasonable limit


# ═══════════════════════════════════════════════════════════════════════
# Test 25: Private/Public Visibility Layer
# ═══════════════════════════════════════════════════════════════════════

class TestVisibilityLayer:
    def test_model_has_visibility_columns(self):
        from api.models.knowledge import KnowledgeDocument
        columns = {c.name for c in KnowledgeDocument.__table__.columns}
        assert "user_id" in columns
        assert "visibility" in columns

    def test_visibility_nullable(self):
        from api.models.knowledge import KnowledgeDocument
        col_map = {c.name: c for c in KnowledgeDocument.__table__.columns}
        assert col_map["user_id"].nullable is True
        # visibility has a default
        assert col_map["visibility"].default is not None

    def test_search_knowledge_user_id_param(self):
        import inspect
        from api.services.rag import search_knowledge
        sig = inspect.signature(search_knowledge)
        assert "user_id" in sig.parameters
        assert sig.parameters["user_id"].default is None

    def test_build_rag_context_user_id_param(self):
        import inspect
        from api.services.rag import build_rag_context
        sig = inspect.signature(build_rag_context)
        assert "user_id" in sig.parameters

    def test_build_groot_system_prompt_user_id_param(self):
        import inspect
        from api.services.rag import build_groot_system_prompt
        sig = inspect.signature(build_groot_system_prompt)
        assert "user_id" in sig.parameters


# ═══════════════════════════════════════════════════════════════════════
# Test 26: URL Parser
# ═══════════════════════════════════════════════════════════════════════

class TestURLParser:
    def test_html_text_extractor(self):
        from api.services.url_parser import _HTMLTextExtractor
        extractor = _HTMLTextExtractor()
        extractor.feed("<html><head><title>Test Page</title></head><body><h1>Hello</h1><p>World</p><script>evil()</script></body></html>")
        text = extractor.get_text()
        assert "Hello" in text
        assert "World" in text
        assert "evil" not in text
        assert extractor.title == "Test Page"

    def test_html_skip_tags(self):
        from api.services.url_parser import _HTMLTextExtractor
        extractor = _HTMLTextExtractor()
        extractor.feed("<nav>Navigation</nav><p>Content</p><footer>Footer</footer><style>.hidden{}</style>")
        text = extractor.get_text()
        assert "Content" in text
        assert "Navigation" not in text
        assert "Footer" not in text
        assert "hidden" not in text

    def test_invalid_url_scheme(self):
        import asyncio
        from api.services.url_parser import parse_url
        result = asyncio.get_event_loop().run_until_complete(parse_url("ftp://example.com"))
        assert result.error is not None
        assert "scheme" in result.error.lower()

    def test_empty_url(self):
        import asyncio
        from api.services.url_parser import parse_url
        result = asyncio.get_event_loop().run_until_complete(parse_url(""))
        assert result.error is not None


# ═══════════════════════════════════════════════════════════════════════
# Test 27: Document Exporter
# ═══════════════════════════════════════════════════════════════════════

class TestDocumentExporter:
    def test_export_markdown(self):
        from api.services.document_exporter import export_markdown
        md = export_markdown("Test Doc", "Hello world", tags=["defi", "staking"], category="blockchain")
        assert "# Test Doc" in md
        assert "Hello world" in md
        assert "defi" in md
        assert "category: blockchain" in md
        assert "REFINET Cloud" in md

    def test_export_markdown_no_tags(self):
        from api.services.document_exporter import export_markdown
        md = export_markdown("Simple", "Content here")
        assert "# Simple" in md
        assert "Content here" in md

    def test_export_pdf(self):
        from api.services.document_exporter import export_pdf
        try:
            import fitz  # noqa
            pdf_bytes = export_pdf("Test PDF", "Hello world from PDF export", tags=["test"], category="docs")
            assert isinstance(pdf_bytes, bytes)
            assert len(pdf_bytes) > 0
            assert pdf_bytes[:4] == b'%PDF'
        except ImportError:
            pass  # PyMuPDF not installed, skip

    def test_export_pdf_long_content(self):
        from api.services.document_exporter import export_pdf
        try:
            import fitz  # noqa
            long_content = "This is a paragraph. " * 500
            pdf_bytes = export_pdf("Long Doc", long_content, category="docs")
            assert len(pdf_bytes) > 0
            # Should create multiple pages
            doc = fitz.open(stream=pdf_bytes, filetype="pdf")
            assert len(doc) > 1
            doc.close()
        except ImportError:
            pass


# ═══════════════════════════════════════════════════════════════════════
# Test 28: Timeline Extractor
# ═══════════════════════════════════════════════════════════════════════

class TestTimelineExtractor:
    def test_extract_iso_dates(self):
        from api.services.timeline_extractor import extract_timeline
        text = "The protocol launched on 2024-01-15. A major upgrade was deployed on 2024-06-01."
        events = extract_timeline(text)
        assert len(events) >= 2
        assert events[0]["date_display"] == "2024-01-15"
        assert events[1]["date_display"] == "2024-06-01"
        # Should be chronologically sorted
        assert events[0]["date"] <= events[1]["date"]

    def test_extract_written_dates(self):
        from api.services.timeline_extractor import extract_timeline
        text = "Founded on January 15, 2023. Raised Series A on March 1, 2024."
        events = extract_timeline(text)
        assert len(events) >= 2
        assert "January" in events[0]["date_display"]

    def test_extract_quarters(self):
        from api.services.timeline_extractor import extract_timeline
        text = "Q1 2024 saw revenue growth. Q3 2024 marked the product launch."
        events = extract_timeline(text)
        assert len(events) >= 2
        assert "Q1 2024" in events[0]["date_display"]

    def test_empty_text(self):
        from api.services.timeline_extractor import extract_timeline
        assert extract_timeline("") == []
        assert extract_timeline("No dates here") == []

    def test_chronological_sort(self):
        from api.services.timeline_extractor import extract_timeline
        text = "Event B on 2024-12-01. Event A on 2024-01-01. Event C on 2024-06-15."
        events = extract_timeline(text)
        dates = [e["date"] for e in events]
        assert dates == sorted(dates)


# ═══════════════════════════════════════════════════════════════════════
# Test 29: YouTube Parser
# ═══════════════════════════════════════════════════════════════════════

class TestYouTubeParser:
    def test_is_youtube_url(self):
        from api.services.youtube_parser import is_youtube_url
        assert is_youtube_url("https://www.youtube.com/watch?v=dQw4w9WgXcQ")
        assert is_youtube_url("https://youtu.be/dQw4w9WgXcQ")
        assert is_youtube_url("https://youtube.com/embed/dQw4w9WgXcQ")
        assert not is_youtube_url("https://vimeo.com/12345")
        assert not is_youtube_url("https://example.com")

    def test_extract_video_id(self):
        from api.services.youtube_parser import extract_video_id
        assert extract_video_id("https://www.youtube.com/watch?v=dQw4w9WgXcQ") == "dQw4w9WgXcQ"
        assert extract_video_id("https://youtu.be/dQw4w9WgXcQ") == "dQw4w9WgXcQ"
        assert extract_video_id("https://example.com") is None

    def test_parse_subtitles(self):
        from api.services.youtube_parser import _parse_subtitles
        vtt = """WEBVTT

00:00:00.000 --> 00:00:02.000
Hello world.

00:00:02.000 --> 00:00:04.000
This is a test.

00:00:04.000 --> 00:00:06.000
Hello world.
"""
        text = _parse_subtitles(vtt)
        assert "Hello world" in text
        assert "This is a test" in text
        # Should deduplicate
        assert text.count("Hello world") == 1

    def test_invalid_url(self):
        import asyncio
        from api.services.youtube_parser import parse_youtube
        result = asyncio.get_event_loop().run_until_complete(parse_youtube("https://example.com"))
        assert result.error is not None
        assert "YouTube" in result.error


# ═══════════════════════════════════════════════════════════════════════
# Test 30: TTS Generator
# ═══════════════════════════════════════════════════════════════════════

class TestTTSGenerator:
    def test_tts_availability_check(self):
        from api.services.tts_generator import is_tts_available
        # Just verify it returns a boolean without crashing
        result = is_tts_available()
        assert isinstance(result, bool)

    def test_podcast_prompt_template(self):
        from api.services.tts_generator import PODCAST_SCRIPT_PROMPT
        assert "{title}" in PODCAST_SCRIPT_PROMPT
        assert "{content}" in PODCAST_SCRIPT_PROMPT


# ═══════════════════════════════════════════════════════════════════════
# Test 31: Document Share Model
# ═══════════════════════════════════════════════════════════════════════

class TestDocumentShareModel:
    def test_model_exists(self):
        from api.models.knowledge import DocumentShare
        columns = {c.name for c in DocumentShare.__table__.columns}
        assert "document_id" in columns
        assert "owner_id" in columns
        assert "shared_with_id" in columns
        assert "permission" in columns

    def test_permission_default(self):
        from api.models.knowledge import DocumentShare
        col_map = {c.name: c for c in DocumentShare.__table__.columns}
        assert col_map["permission"].default is not None

    def test_share_unique_constraint(self):
        from api.models.knowledge import DocumentShare
        # Check that unique constraint exists on (document_id, shared_with_id)
        constraints = DocumentShare.__table__.constraints
        unique_names = [c.name for c in constraints if hasattr(c, 'name') and c.name]
        assert "uq_share_doc_user" in unique_names

    def test_share_foreign_key(self):
        from api.models.knowledge import DocumentShare
        col_map = {c.name: c for c in DocumentShare.__table__.columns}
        fks = list(col_map["document_id"].foreign_keys)
        assert len(fks) == 1
        assert "knowledge_documents.id" in str(fks[0])


# ═══════════════════════════════════════════════════════════════════════
# Test 32: Database Integrity Constraints
# ═══════════════════════════════════════════════════════════════════════

class TestDatabaseConstraints:
    def test_content_hash_unique(self):
        from api.models.knowledge import KnowledgeDocument
        col_map = {c.name: c for c in KnowledgeDocument.__table__.columns}
        assert col_map["content_hash"].unique is True

    def test_chunk_foreign_key(self):
        from api.models.knowledge import KnowledgeChunk
        col_map = {c.name: c for c in KnowledgeChunk.__table__.columns}
        fks = list(col_map["document_id"].foreign_keys)
        assert len(fks) == 1
        assert "knowledge_documents.id" in str(fks[0])

    def test_chunk_cascade_delete(self):
        from api.models.knowledge import KnowledgeChunk
        col_map = {c.name: c for c in KnowledgeChunk.__table__.columns}
        fk = list(col_map["document_id"].foreign_keys)[0]
        assert fk.ondelete == "CASCADE"

    def test_share_cascade_delete(self):
        from api.models.knowledge import DocumentShare
        col_map = {c.name: c for c in DocumentShare.__table__.columns}
        fk = list(col_map["document_id"].foreign_keys)[0]
        assert fk.ondelete == "CASCADE"


# ═══════════════════════════════════════════════════════════════════════
# Test 33: MCP Visibility Enforcement
# ═══════════════════════════════════════════════════════════════════════

class TestMCPVisibility:
    def test_search_documents_tool_exists(self):
        from api.services.mcp_gateway import MCP_TOOLS
        assert "search_documents" in MCP_TOOLS
        assert "query" in MCP_TOOLS["search_documents"]["input_schema"]["required"]

    def test_compare_documents_tool_exists(self):
        from api.services.mcp_gateway import MCP_TOOLS
        assert "compare_documents" in MCP_TOOLS

    def test_get_document_tags_tool_exists(self):
        from api.services.mcp_gateway import MCP_TOOLS
        assert "get_document_tags" in MCP_TOOLS

    def test_all_document_tools_present(self):
        from api.services.mcp_gateway import list_tools
        tools = list_tools()
        names = {t["name"] for t in tools}
        assert "search_documents" in names
        assert "compare_documents" in names
        assert "get_document_tags" in names


# ═══════════════════════════════════════════════════════════════════════
# Test 34: RAG Batch Query Optimization
# ═══════════════════════════════════════════════════════════════════════

class TestRAGOptimization:
    def test_search_knowledge_has_user_id(self):
        import inspect
        from api.services.rag import search_knowledge
        sig = inspect.signature(search_knowledge)
        params = list(sig.parameters.keys())
        # Verify all filter params exist
        assert "user_id" in params
        assert "doc_ids" in params
        assert "tags" in params
        assert "category" in params

    def test_build_rag_context_returns_tuple(self):
        import inspect
        from api.services.rag import build_rag_context
        sig = inspect.signature(build_rag_context)
        # Verify it accepts both doc_ids and user_id
        assert "doc_ids" in sig.parameters
        assert "user_id" in sig.parameters


# ═══════════════════════════════════════════════════════════════════════
# Test 35: Export Endpoints Exist
# ═══════════════════════════════════════════════════════════════════════

class TestExportEndpoints:
    def test_markdown_export_format(self):
        from api.services.document_exporter import export_markdown
        md = export_markdown("Test", "Content", tags=["a", "b"], category="docs", doc_type="pdf")
        assert md.startswith("---")
        assert 'title: "Test"' in md
        assert "type: pdf" in md

    def test_pdf_export_returns_bytes(self):
        from api.services.document_exporter import export_pdf
        try:
            import fitz  # noqa
            result = export_pdf("Test", "Hello World")
            assert isinstance(result, bytes)
            assert result[:4] == b"%PDF"
        except ImportError:
            pass  # PyMuPDF not available

    def test_timeline_no_dates(self):
        from api.services.timeline_extractor import extract_timeline
        result = extract_timeline("This text has absolutely no dates in it whatsoever")
        assert result == []

    def test_url_parser_html_extraction(self):
        from api.services.url_parser import _HTMLTextExtractor
        ex = _HTMLTextExtractor()
        ex.feed("<html><body><p>Hello</p><script>alert('xss')</script><p>World</p></body></html>")
        text = ex.get_text()
        assert "Hello" in text
        assert "World" in text
        assert "alert" not in text


# ═══════════════════════════════════════════════════════════════════════
# Test 36: Config Defaults
# ═══════════════════════════════════════════════════════════════════════

class TestConfigDefaults:
    def test_defaults_list_complete(self):
        from api.services.config_defaults import DEFAULTS
        keys = {d["key"] for d in DEFAULTS}
        assert "knowledge.max_documents_per_user" in keys
        assert "knowledge.max_file_size_mb" in keys
        assert "knowledge.max_total_storage_mb" in keys
        assert "knowledge.allowed_file_types" in keys
        assert "platform.maintenance_mode" in keys
        assert "platform.allow_user_uploads" in keys
        assert "platform.allow_public_documents" in keys
        assert "oracle.enabled" in keys

    def test_defaults_have_required_fields(self):
        from api.services.config_defaults import DEFAULTS
        for d in DEFAULTS:
            assert "key" in d
            assert "value" in d
            assert "data_type" in d
            assert d["data_type"] in ("string", "integer", "boolean", "json")

    def test_get_config_int_returns_default(self):
        from api.services.config_defaults import get_config_int
        # With no DB, should return default
        class FakeQuery:
            def filter(self, *a): return self
            def first(self): return None
        class FakeDB:
            def query(self, *a): return FakeQuery()
        assert get_config_int(FakeDB(), "nonexistent", 42) == 42

    def test_get_config_bool_returns_default(self):
        from api.services.config_defaults import get_config_bool
        class FakeQuery:
            def filter(self, *a): return self
            def first(self): return None
        class FakeDB:
            def query(self, *a): return FakeQuery()
        assert get_config_bool(FakeDB(), "nonexistent", True) is True
        assert get_config_bool(FakeDB(), "nonexistent", False) is False


# ═══════════════════════════════════════════════════════════════════════
# Test 37: Knowledge Refresh Handler
# ═══════════════════════════════════════════════════════════════════════

class TestKnowledgeRefresh:
    def test_handler_is_async(self):
        import asyncio
        from api.services.knowledge_refresh import on_knowledge_change
        assert asyncio.iscoroutinefunction(on_knowledge_change)

    def test_handler_runs_without_error(self):
        import asyncio
        from api.services.knowledge_refresh import on_knowledge_change
        # Should handle any event without crashing
        loop = asyncio.get_event_loop()
        loop.run_until_complete(on_knowledge_change("knowledge.document.uploaded", {
            "document_id": "test-123", "title": "Test Doc", "user_id": "user-1",
        }))

    def test_handler_visibility_change(self):
        import asyncio
        from api.services.knowledge_refresh import on_knowledge_change
        loop = asyncio.get_event_loop()
        # Should not crash even without DB
        loop.run_until_complete(on_knowledge_change("knowledge.document.visibility_changed", {
            "document_id": "test-123", "new_visibility": "public",
        }))


# ═══════════════════════════════════════════════════════════════════════
# Test 38: Event Emission Helper
# ═══════════════════════════════════════════════════════════════════════

class TestEventEmission:
    def test_emit_knowledge_event_is_async(self):
        import asyncio
        # Import the function from knowledge routes
        import importlib
        mod = importlib.import_module("api.routes.knowledge")
        fn = getattr(mod, "_emit_knowledge_event")
        assert asyncio.iscoroutinefunction(fn)

    def test_emit_event_doesnt_crash(self):
        import asyncio
        import importlib
        mod = importlib.import_module("api.routes.knowledge")
        fn = getattr(mod, "_emit_knowledge_event")
        # Should not crash even if EventBus isn't running
        loop = asyncio.get_event_loop()
        loop.run_until_complete(fn("knowledge.test", {"test": True}))

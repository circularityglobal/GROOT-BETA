"""
REFINET Cloud — Sovereign Document Parser
Extracts plain text + metadata from PDF, DOCX, XLSX, CSV, TXT, MD, JSON, Solidity.
Zero external API calls. All MIT/BSD licensed dependencies.
Plugs directly into the existing knowledge_chunks pipeline.
"""

import csv
import io
import json
import logging
import os
import re
from dataclasses import dataclass, field
from typing import Optional

logger = logging.getLogger("refinet.document_parser")

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".csv", ".txt", ".md", ".json", ".sol"}

EXT_TO_DOCTYPE = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".xlsx": "xlsx",
    ".csv": "csv",
    ".txt": "txt",
    ".md": "md",
    ".json": "json",
    ".sol": "sol",
}


@dataclass
class ParseResult:
    """Output of document parsing."""
    text: str                                      # Extracted plain text
    doc_type: str                                  # File type without dot
    page_count: Optional[int] = None               # Pages/sheets if applicable
    metadata: dict = field(default_factory=dict)   # Author, dates, file_size, etc.
    error: Optional[str] = None                    # Parse error if partial extraction


def parse_file(file_bytes: bytes, filename: str) -> ParseResult:
    """
    Main entry point. Dispatches to format-specific parser based on extension.
    Returns ParseResult with extracted text and metadata.
    """
    ext = os.path.splitext(filename)[1].lower()
    if ext not in SUPPORTED_EXTENSIONS:
        return ParseResult(
            text="",
            doc_type=ext.lstrip(".") or "unknown",
            error=f"Unsupported file type: {ext}. Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}",
        )

    doc_type = EXT_TO_DOCTYPE.get(ext, ext.lstrip("."))
    base_metadata = {"file_size": len(file_bytes), "filename": filename}

    parsers = {
        ".pdf": parse_pdf,
        ".docx": _parse_docx,
        ".xlsx": _parse_xlsx,
        ".csv": _parse_csv,
        ".json": _parse_json,
        ".sol": _parse_solidity,
        ".txt": _parse_text,
        ".md": _parse_text,
    }

    parser = parsers.get(ext, _parse_text)
    try:
        result = parser(file_bytes, filename)
        result.metadata.update(base_metadata)
        return result
    except Exception as e:
        logger.error(f"Parse error for {filename}: {e}")
        # Attempt raw text fallback
        fallback_text = ""
        try:
            fallback_text = file_bytes.decode("utf-8", errors="replace")
        except Exception:
            pass
        return ParseResult(
            text=fallback_text,
            doc_type=doc_type,
            metadata=base_metadata,
            error=f"Parse error: {str(e)}",
        )


# ── PDF ───────────────────────────────────────────────────────────────

def parse_pdf(file_bytes: bytes, filename: str) -> ParseResult:
    """Extract text from PDF using PyMuPDF (fitz). Page-by-page extraction."""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        return ParseResult(
            text="",
            doc_type="pdf",
            error="PyMuPDF not installed. Run: pip install PyMuPDF",
        )

    doc = fitz.open(stream=file_bytes, filetype="pdf")
    pages = []
    for i in range(len(doc)):
        page = doc.load_page(i)
        text = page.get_text("text")
        if text.strip():
            pages.append(text)

    # Extract PDF metadata
    meta = doc.metadata or {}
    metadata = {}
    if meta.get("author"):
        metadata["author"] = meta["author"]
    if meta.get("title"):
        metadata["pdf_title"] = meta["title"]
    if meta.get("subject"):
        metadata["subject"] = meta["subject"]
    if meta.get("creationDate"):
        metadata["created"] = meta["creationDate"]
    if meta.get("modDate"):
        metadata["modified"] = meta["modDate"]
    if meta.get("producer"):
        metadata["producer"] = meta["producer"]

    page_count = len(doc)
    doc.close()

    full_text = "\n\n".join(pages)

    return ParseResult(
        text=full_text,
        doc_type="pdf",
        page_count=page_count,
        metadata=metadata,
    )


# ── DOCX ──────────────────────────────────────────────────────────────

def _parse_docx(file_bytes: bytes, filename: str) -> ParseResult:
    """Extract text from DOCX using python-docx. Paragraph-by-paragraph."""
    try:
        import docx
    except ImportError:
        return ParseResult(
            text="",
            doc_type="docx",
            error="python-docx not installed. Run: pip install python-docx",
        )

    doc = docx.Document(io.BytesIO(file_bytes))

    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append("\t".join(cells))

    # Extract metadata from core properties
    metadata = {}
    try:
        props = doc.core_properties
        if props.author:
            metadata["author"] = props.author
        if props.title:
            metadata["docx_title"] = props.title
        if props.subject:
            metadata["subject"] = props.subject
        if props.created:
            metadata["created"] = str(props.created)
        if props.modified:
            metadata["modified"] = str(props.modified)
    except Exception:
        pass

    full_text = "\n\n".join(paragraphs)

    return ParseResult(
        text=full_text,
        doc_type="docx",
        page_count=None,  # DOCX doesn't reliably expose page count
        metadata=metadata,
    )


# ── XLSX ──────────────────────────────────────────────────────────────

def _parse_xlsx(file_bytes: bytes, filename: str) -> ParseResult:
    """Extract text from XLSX using openpyxl. Sheet-by-sheet, row-by-row."""
    try:
        import openpyxl
    except ImportError:
        return ParseResult(
            text="",
            doc_type="xlsx",
            error="openpyxl not installed. Run: pip install openpyxl",
        )

    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
    sections = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            line = "\t".join(cells).strip()
            if line:
                rows.append(line)
        if rows:
            sections.append(f"=== Sheet: {sheet_name} ===\n" + "\n".join(rows))

    sheet_names = list(wb.sheetnames)
    sheet_count = len(sheet_names)
    wb.close()

    full_text = "\n\n".join(sections)

    return ParseResult(
        text=full_text,
        doc_type="xlsx",
        page_count=sheet_count,
        metadata={"sheet_names": sheet_names},
    )


# ── CSV ───────────────────────────────────────────────────────────────

def _parse_csv(file_bytes: bytes, filename: str) -> ParseResult:
    """Extract text from CSV using stdlib csv module."""
    # Try UTF-8 first, then UTF-8-BOM, then latin-1
    text_content = None
    for encoding in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            text_content = file_bytes.decode(encoding)
            break
        except UnicodeDecodeError:
            continue

    if text_content is None:
        return ParseResult(text="", doc_type="csv", error="Could not decode CSV file")

    reader = csv.reader(io.StringIO(text_content))
    rows = []
    row_count = 0
    for row in reader:
        line = "\t".join(row).strip()
        if line:
            rows.append(line)
            row_count += 1

    full_text = "\n".join(rows)

    return ParseResult(
        text=full_text,
        doc_type="csv",
        metadata={"row_count": row_count},
    )


# ── JSON ──────────────────────────────────────────────────────────────

def _parse_json(file_bytes: bytes, filename: str) -> ParseResult:
    """Extract text from JSON. Handles ABI JSON specially."""
    text_content = file_bytes.decode("utf-8", errors="replace")
    data = json.loads(text_content)

    # Special case: ABI JSON (list of objects with "type" and optionally "name")
    if isinstance(data, list) and len(data) > 0 and isinstance(data[0], dict):
        if "type" in data[0] and data[0].get("type") in ("function", "event", "constructor", "fallback", "receive", "error"):
            return _parse_abi_json(data, filename)

    # General JSON: pretty-print
    pretty = json.dumps(data, indent=2, ensure_ascii=False)

    return ParseResult(
        text=pretty,
        doc_type="json",
        metadata={"json_type": type(data).__name__},
    )


def _parse_abi_json(abi_data: list, filename: str) -> ParseResult:
    """Convert ABI JSON to human-readable contract interface summary."""
    lines = [f"Smart Contract ABI: {filename}", ""]

    functions = []
    events = []
    for item in abi_data:
        item_type = item.get("type", "")
        name = item.get("name", "")
        inputs = item.get("inputs", [])
        outputs = item.get("outputs", [])

        if item_type in ("function", "constructor", "fallback", "receive"):
            input_sig = ", ".join(f"{i.get('type', '?')} {i.get('name', '')}" for i in inputs)
            output_sig = ", ".join(f"{o.get('type', '?')}" for o in outputs)
            mutability = item.get("stateMutability", "")
            line = f"  {name}({input_sig})"
            if output_sig:
                line += f" → ({output_sig})"
            if mutability:
                line += f"  [{mutability}]"
            functions.append(line)

        elif item_type == "event":
            input_sig = ", ".join(
                f"{'indexed ' if i.get('indexed') else ''}{i.get('type', '?')} {i.get('name', '')}"
                for i in inputs
            )
            events.append(f"  {name}({input_sig})")

    if functions:
        lines.append(f"Functions ({len(functions)}):")
        lines.extend(functions)
        lines.append("")

    if events:
        lines.append(f"Events ({len(events)}):")
        lines.extend(events)

    return ParseResult(
        text="\n".join(lines),
        doc_type="json",
        metadata={"is_abi": True, "function_count": len(functions), "event_count": len(events)},
    )


# ── Solidity ──────────────────────────────────────────────────────────

def _parse_solidity(file_bytes: bytes, filename: str) -> ParseResult:
    """Extract text from Solidity files. Preserve NatSpec comments as documentation."""
    text = file_bytes.decode("utf-8", errors="replace")

    # Extract NatSpec comments
    natspec_pattern = r'///\s*(.*?)$|/\*\*(.*?)\*/'
    natspec_matches = re.findall(natspec_pattern, text, re.MULTILINE | re.DOTALL)
    natspec_docs = []
    for single, multi in natspec_matches:
        doc = (single or multi).strip()
        if doc:
            # Clean up multi-line NatSpec
            doc = re.sub(r'\n\s*\*\s*', '\n', doc)
            natspec_docs.append(doc)

    # Extract contract/interface names
    contract_names = re.findall(r'\b(?:contract|interface|library|abstract\s+contract)\s+(\w+)', text)

    # Extract pragma
    pragma_match = re.search(r'pragma\s+solidity\s+([^;]+);', text)
    pragma = pragma_match.group(1).strip() if pragma_match else None

    metadata = {}
    if contract_names:
        metadata["contracts"] = contract_names
    if pragma:
        metadata["solidity_version"] = pragma
    if natspec_docs:
        metadata["natspec_count"] = len(natspec_docs)

    return ParseResult(
        text=text,
        doc_type="sol",
        metadata=metadata,
    )


# ── Text / Markdown ──────────────────────────────────────────────────

def _parse_text(file_bytes: bytes, filename: str) -> ParseResult:
    """Plain text / Markdown passthrough with encoding detection."""
    # Try UTF-8 first, then latin-1 fallback
    text = None
    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            text = file_bytes.decode(encoding)
            break
        except UnicodeDecodeError:
            continue

    if text is None:
        text = file_bytes.decode("utf-8", errors="replace")

    ext = os.path.splitext(filename)[1].lower()
    doc_type = "md" if ext == ".md" else "txt"

    metadata = {}
    if doc_type == "md":
        # Extract headings from markdown
        headings = re.findall(r'^#{1,3}\s+(.+)$', text, re.MULTILINE)
        if headings:
            metadata["headings"] = headings[:20]  # Cap at 20

    return ParseResult(
        text=text,
        doc_type=doc_type,
        metadata=metadata,
    )
